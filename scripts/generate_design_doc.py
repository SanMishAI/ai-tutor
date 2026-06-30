"""
Converts DESIGN.md to a branded DESIGN.docx.
Run from the project root: python scripts/generate_design_doc.py
Deletes any existing DESIGN.docx before writing a fresh copy.
"""

import re
import sys
import io
import tempfile
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

try:
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import matplotlib.patches as mpatches
    from matplotlib.patches import FancyArrowPatch, FancyBboxPatch
    HAS_MPL = True
except ImportError:
    HAS_MPL = False
    print("WARNING: matplotlib not installed — diagrams will be skipped. Run: pip install matplotlib")

ROOT     = Path(__file__).parent.parent
SRC      = ROOT / "DESIGN.md"
DEST     = ROOT / "DESIGN.docx"
LOGO_PNG = ROOT / "scripts" / "logo.png"

# ── Brand palette ─────────────────────────────────────────────────────────────
NAVY      = RGBColor(0x00, 0x09, 0x36)
GOLD      = RGBColor(0xFD, 0xC8, 0x00)
ORANGE    = RGBColor(0xE3, 0x4C, 0x00)
SKY       = RGBColor(0x56, 0xDB, 0xFF)
BLUE      = RGBColor(0x00, 0x66, 0xCB)
TAG_BLUE  = RGBColor(0x1D, 0xA4, 0xF3)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
DARK_GRAY = RGBColor(0x1E, 0x29, 0x3B)
MID_GRAY  = RGBColor(0x47, 0x55, 0x69)
CODE_FG   = RGBColor(0x0F, 0x17, 0x2A)


# ── XML helpers ───────────────────────────────────────────────────────────────
def _rgb_hex(color: RGBColor) -> str:
    return f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"


def set_cell_shd(cell, fill_hex: str):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex.replace("#", ""))
    tcPr.append(shd)


def set_cell_border(cell, color="E2E8F0"):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(old)
    tcB = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color.replace("#", ""))
        tcB.append(el)
    tcPr.append(tcB)


def _insert_pBdr(pPr, side: str, color: str, sz: str = "6", space: str = "4"):
    """Insert w:pBdr into pPr BEFORE w:spacing/w:jc to satisfy OOXML schema order."""
    for old in pPr.findall(qn("w:pBdr")):
        pPr.remove(old)
    pBdr = OxmlElement("w:pBdr")
    el = OxmlElement(f"w:{side}")
    el.set(qn("w:val"),   "single")
    el.set(qn("w:sz"),    sz)
    el.set(qn("w:space"), space)
    el.set(qn("w:color"), color.replace("#", ""))
    pBdr.append(el)
    # Insert before w:spacing, w:ind, or w:jc — whichever comes first
    for tag in ("w:spacing", "w:ind", "w:jc", "w:rPr"):
        existing = pPr.find(qn(tag))
        if existing is not None:
            existing.addprevious(pBdr)
            return
    pPr.append(pBdr)


def para_spacing(para, before_twips=0, after_twips=0):
    pPr = para._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:spacing")):
        pPr.remove(old)
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:before"), str(before_twips))
    sp.set(qn("w:after"),  str(after_twips))
    pPr.append(sp)


def set_para_shd(para, fill_hex: str):
    pPr = para._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:shd")):
        pPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex.replace("#", ""))
    # Insert before w:spacing/w:jc to maintain OOXML schema order
    for tag in ("w:spacing", "w:ind", "w:jc"):
        existing = pPr.find(qn(tag))
        if existing is not None:
            existing.addprevious(shd)
            return
    pPr.append(shd)


def add_field_run(para_elem, field_code: str):
    """Append a properly-structured field (begin/instrText/end) to a paragraph element."""
    for fc_type, content in [("begin", None), (None, field_code), ("end", None)]:
        r = OxmlElement("w:r")
        if fc_type is not None:
            fc = OxmlElement("w:fldChar")
            fc.set(qn("w:fldCharType"), fc_type)
            r.append(fc)
        else:
            it = OxmlElement("w:instrText")
            it.set(qn("xml:space"), "preserve")
            it.text = f" {content} "
            r.append(it)
        para_elem.append(r)


def xml_text_run(para_elem, text: str, color: RGBColor, size_pt: float = 8.0, bold=False):
    """Append a plain text run directly to a paragraph XML element."""
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    if bold:
        rPr.append(OxmlElement("w:b"))
    fn = OxmlElement("w:rFonts")
    fn.set(qn("w:ascii"), "Calibri")
    fn.set(qn("w:hAnsi"), "Calibri")
    rPr.append(fn)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size_pt * 2)))
    rPr.append(sz)
    clr = OxmlElement("w:color")
    clr.set(qn("w:val"), _rgb_hex(color))
    rPr.append(clr)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    para_elem.append(r)


# ── Header / footer ───────────────────────────────────────────────────────────
def add_header_footer(doc: Document):
    section = doc.sections[0]

    # ── Header ────────────────────────────────────────────────────────────────
    header = section.header
    header.is_linked_to_previous = False
    for p in header.paragraphs:
        p.clear()

    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Add pBdr FIRST, then spacing — satisfies OOXML element order
    pPr = hp._p.get_or_add_pPr()
    _insert_pBdr(pPr, "bottom", "FDC800", sz="6", space="4")
    para_spacing(hp, 0, 60)

    if LOGO_PNG.exists():
        logo_run = hp.add_run()
        logo_run.add_picture(str(LOGO_PNG), width=Inches(0.28))

    r_sel = hp.add_run("  Select")
    r_sel.font.name = "Calibri"; r_sel.font.bold = True
    r_sel.font.size = Pt(10);    r_sel.font.color.rgb = SKY

    r_ed = hp.add_run("Ed")
    r_ed.font.name = "Calibri"; r_ed.font.bold = True
    r_ed.font.size = Pt(10);    r_ed.font.color.rgb = GOLD

    r_title = hp.add_run("  -  Design & Architecture Document")
    r_title.font.name = "Calibri"
    r_title.font.size = Pt(8.5)
    r_title.font.color.rgb = MID_GRAY

    # ── Footer ────────────────────────────────────────────────────────────────
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p.clear()

    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add top border BEFORE spacing
    fPr = fp._p.get_or_add_pPr()
    _insert_pBdr(fPr, "top", "E2E8F0", sz="4", space="4")
    para_spacing(fp, 60, 0)

    xml_text_run(fp._p, "SelectEd  -  Confidential  -  Page ", MID_GRAY, 8.0)
    add_field_run(fp._p, "PAGE")
    xml_text_run(fp._p, " of ", MID_GRAY, 8.0)
    add_field_run(fp._p, "NUMPAGES")


# ── Inline markdown renderer ──────────────────────────────────────────────────
def render_inline(para, text: str, base_color=None, base_bold=False, base_size=None):
    pattern = re.compile(r"(\*\*.*?\*\*|\*[^*].*?[^*]\*|`[^`]+`)")
    for part in pattern.split(text):
        if not part:
            continue
        run = para.add_run()
        if part.startswith("**") and part.endswith("**"):
            run.text = part[2:-2]
            run.bold = True
            if base_color:
                run.font.color.rgb = base_color
        elif part.startswith("`") and part.endswith("`"):
            run.text           = part[1:-1]
            run.font.name      = "Courier New"
            run.font.size      = Pt(9)
            run.font.color.rgb = CODE_FG
        elif part.startswith("*") and part.endswith("*"):
            run.text   = part[1:-1]
            run.italic = True
            if base_color:
                run.font.color.rgb = base_color
        else:
            run.text = part
            if base_color:
                run.font.color.rgb = base_color
            if base_bold:
                run.bold = True
            if base_size:
                run.font.size = Pt(base_size)


def parse_md_table(lines):
    headers, rows = [], []
    for idx, line in enumerate(lines):
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        if idx == 0:
            headers = cells
        elif re.match(r"^[\s|:\-]+$", line):
            continue
        else:
            rows.append(cells)
    return headers, rows


# ── Document builder ──────────────────────────────────────────────────────────
def build_doc(md_text: str) -> Document:
    doc = Document()

    sec = doc.sections[0]
    sec.page_width    = Inches(8.5)
    sec.page_height   = Inches(11)
    sec.left_margin   = Inches(1.1)
    sec.right_margin  = Inches(1.1)
    sec.top_margin    = Inches(1.1)
    sec.bottom_margin = Inches(0.9)

    # ── Styles ────────────────────────────────────────────────────────────────
    styles = doc.styles

    def ensure_style(name, base="Normal"):
        if name in [s.name for s in styles]:
            return styles[name]
        st = styles.add_style(name, 1)
        st.base_style = styles[base]
        return st

    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = DARK_GRAY

    h1 = styles["Heading 1"]
    h1.font.name = "Calibri"; h1.font.size = Pt(20); h1.font.bold = True
    h1.font.color.rgb = NAVY
    h1.paragraph_format.space_before = Pt(20)
    h1.paragraph_format.space_after  = Pt(6)

    h2 = styles["Heading 2"]
    h2.font.name = "Calibri"; h2.font.size = Pt(14); h2.font.bold = True
    h2.font.color.rgb = BLUE
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after  = Pt(4)

    h3 = styles["Heading 3"]
    h3.font.name = "Calibri"; h3.font.size = Pt(11.5); h3.font.bold = True
    h3.font.color.rgb = MID_GRAY
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after  = Pt(2)

    code_st = ensure_style("MD Code", "Normal")
    code_st.font.name = "Courier New"
    code_st.font.size = Pt(8.5)
    code_st.font.color.rgb = CODE_FG
    code_st.paragraph_format.left_indent  = Cm(0.6)
    code_st.paragraph_format.space_before = Pt(0)
    code_st.paragraph_format.space_after  = Pt(0)

    ensure_style("MD Note", "Normal").font.italic = True

    # ── Cover page ────────────────────────────────────────────────────────────
    if LOGO_PNG.exists():
        lp = doc.add_paragraph()
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para_spacing(lp, 720, 100)
        lp.add_run().add_picture(str(LOGO_PNG), width=Inches(2.0))

    wm = doc.add_paragraph()
    wm.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_spacing(wm, 0, 30)
    r_s = wm.add_run("Select")
    r_s.font.name = "Calibri"; r_s.font.size = Pt(42); r_s.font.bold = True; r_s.font.color.rgb = SKY
    r_e = wm.add_run("Ed")
    r_e.font.name = "Calibri"; r_e.font.size = Pt(42); r_e.font.bold = True; r_e.font.color.rgb = GOLD

    tag = doc.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_spacing(tag, 0, 30)
    for txt, clr in [("Sharpen", GOLD), ("  -  ", MID_GRAY), ("Sit", TAG_BLUE), ("  -  ", MID_GRAY), ("Succeed.", ORANGE)]:
        r = tag.add_run(txt)
        r.font.name = "Calibri"; r.font.size = Pt(13)
        r.font.bold = txt not in ("  -  ",)
        r.font.color.rgb = clr

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_spacing(sub, 20, 50)
    rs = sub.add_run("AI-powered exam preparation for Australian & international competitions")
    rs.font.name = "Calibri"; rs.font.size = Pt(12); rs.font.italic = True; rs.font.color.rgb = MID_GRAY

    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_spacing(rule, 0, 30)
    rr = rule.add_run("---" * 20)
    rr.font.color.rgb = GOLD; rr.font.size = Pt(10)

    dt = doc.add_paragraph()
    dt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_spacing(dt, 0, 20)
    rd = dt.add_run("Design & Architecture Document")
    rd.font.name = "Calibri"; rd.font.size = Pt(16); rd.font.bold = True; rd.font.color.rgb = NAVY

    mt = doc.add_paragraph()
    mt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_spacing(mt, 0, 720)
    rm = mt.add_run("Santrupta Mishra (San)  -  Founder, SelectEd  -  June 2026")
    rm.font.name = "Calibri"; rm.font.size = Pt(10); rm.font.color.rgb = MID_GRAY

    doc.add_page_break()
    add_header_footer(doc)

    # ── Parse markdown ────────────────────────────────────────────────────────
    lines   = md_text.splitlines()
    i       = 0
    in_code = False
    code_buf: list[str] = []

    while i < len(lines):
        line = lines[i]

        # Code fence
        if line.strip().startswith("```"):
            if not in_code:
                in_code  = True
                code_buf = []
            else:
                in_code = False
                for cl in code_buf:
                    cp = doc.add_paragraph(style="MD Code")
                    para_spacing(cp, 0, 0)
                    cp.add_run(cl if cl else " ")
                    set_para_shd(cp, "F1F5F9")
                doc.add_paragraph()
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^[-_]{3,}$", line.strip()):
            hr = doc.add_paragraph()
            para_spacing(hr, 60, 60)
            hr.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rr2 = hr.add_run("-" * 60)
            rr2.font.color.rgb = GOLD; rr2.font.size = Pt(9)
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,3})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            text  = m.group(2).strip()
            smap  = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}
            para  = doc.add_paragraph(style=smap.get(level, "Heading 3"))
            render_inline(para, text)
            i += 1
            continue

        # Table
        if line.strip().startswith("|"):
            tbl_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                tbl_lines.append(lines[i])
                i += 1
            headers, rows = parse_md_table(tbl_lines)
            if not headers:
                continue
            ncols = len(headers)
            tbl   = doc.add_table(rows=1 + len(rows), cols=ncols)
            tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
            tbl.style = "Table Grid"

            for ci, h in enumerate(headers):
                cell = tbl.rows[0].cells[ci]
                cell.text = ""
                p = cell.paragraphs[0]
                r = p.add_run(h)
                r.bold = True; r.font.name = "Calibri"
                r.font.size = Pt(9.5); r.font.color.rgb = WHITE
                set_cell_shd(cell, "000936")
                set_cell_border(cell, "000936")

            for ri, row in enumerate(rows):
                bg = "F8FAFC" if ri % 2 == 0 else "FFFFFF"
                for ci, cell_text in enumerate(row):
                    if ci >= ncols:
                        break
                    cell = tbl.rows[ri + 1].cells[ci]
                    cell.text = ""
                    p = cell.paragraphs[0]
                    para_spacing(p, 30, 30)
                    render_inline(p, cell_text.strip(), base_color=DARK_GRAY, base_size=9)
                    set_cell_shd(cell, bg)
                    set_cell_border(cell, "E2E8F0")

            doc.add_paragraph()
            continue

        # Bullet / numbered list
        m_b = re.match(r"^(\s*)([-*]|\d+\.)\s+(.*)", line)
        if m_b:
            indent = len(m_b.group(1))
            is_num = bool(re.match(r"\d+\.", m_b.group(2)))
            bp = doc.add_paragraph(style="List Number" if is_num else "List Bullet")
            bp.paragraph_format.left_indent = Cm(0.8 + indent * 0.4)
            render_inline(bp, m_b.group(3), base_color=DARK_GRAY)
            i += 1
            continue

        # Blank
        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        # Normal paragraph
        pp = doc.add_paragraph(style="Normal")
        para_spacing(pp, 2, 4)
        render_inline(pp, stripped, base_color=DARK_GRAY)
        i += 1

    return doc


# ── Diagram generators ────────────────────────────────────────────────────────

def _save_fig_to_bytes(fig) -> bytes:
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf.read()


def make_architecture_diagram() -> bytes | None:
    if not HAS_MPL:
        return None
    fig, ax = plt.subplots(figsize=(10, 5))
    ax.set_xlim(0, 10); ax.set_ylim(0, 5); ax.axis("off")
    fig.patch.set_facecolor("#F8FAFC")
    ax.set_facecolor("#F8FAFC")

    # colour palette
    C_NAVY = "#000936"; C_GOLD = "#FDC800"; C_BLUE = "#0066CB"
    C_ORANGE = "#E34C00"; C_TEAL = "#0891B2"; C_GREEN = "#059669"
    C_PURPLE = "#7C3AED"; C_LIGHT = "#EFF6FF"; C_GRAY = "#F1F5F9"

    def box(x, y, w, h, label, sub="", fill=C_LIGHT, edge=C_BLUE, lc=C_NAVY):
        rect = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.08",
                               facecolor=fill, edgecolor=edge, linewidth=1.8, zorder=2)
        ax.add_patch(rect)
        ax.text(x + w/2, y + h/2 + (0.12 if sub else 0), label,
                ha="center", va="center", fontsize=8.5, fontweight="bold", color=lc, zorder=3)
        if sub:
            ax.text(x + w/2, y + h/2 - 0.22, sub,
                    ha="center", va="center", fontsize=6.5, color="#64748B", zorder=3)

    def arrow(x1, y1, x2, y2, color=C_BLUE):
        ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle="-|>", color=color, lw=1.6),
                    zorder=1)

    # Title
    ax.text(5, 4.7, "SelectEd — System Architecture", ha="center", va="top",
            fontsize=12, fontweight="bold", color=C_NAVY)

    # Client layer
    box(0.2, 3.2, 2.0, 1.0, "Browser / App", "React 19 + Tailwind v4", fill="#EFF6FF", edge=C_BLUE)

    # Next.js layer
    box(3.2, 3.2, 2.0, 1.0, "Next.js 16", "App Router + proxy.ts", fill="#F0FDF4", edge=C_GREEN)

    # Auth
    box(5.8, 3.8, 1.6, 0.7, "Clerk v7", "Auth / JWT", fill="#FAF5FF", edge=C_PURPLE)

    # AI
    box(5.8, 2.8, 1.6, 0.7, "Claude Sonnet", "Anthropic API", fill="#FFF7ED", edge=C_ORANGE)

    # Stripe
    box(5.8, 1.8, 1.6, 0.7, "Stripe", "Billing / Webhook", fill="#FFF1F2", edge="#E11D48")

    # DB
    box(7.8, 3.2, 1.8, 1.0, "Neon PostgreSQL", "Prisma 7 ORM\nAWS ap-southeast-2", fill=C_GRAY, edge=C_TEAL, lc=C_TEAL)

    # Redis
    box(7.8, 1.9, 1.8, 0.9, "Upstash Redis", "Rate limiting\nSliding window", fill=C_GRAY, edge="#DC2626", lc="#DC2626")

    # Vercel layer (background)
    vercel_bg = FancyBboxPatch((2.8, 2.9), 2.8, 1.5,
                                boxstyle="round,pad=0.1",
                                facecolor="#F0FDF4", edgecolor=C_GREEN,
                                linewidth=1.2, linestyle="--", zorder=1, alpha=0.4)
    ax.add_patch(vercel_bg)
    ax.text(4.2, 4.55, "Vercel Edge Network", ha="center", fontsize=6.5, color=C_GREEN, style="italic")

    # Arrows
    arrow(2.2, 3.7, 3.2, 3.7)              # browser → Next.js
    arrow(5.2, 3.85, 5.8, 3.85+0.12)       # Next.js → Clerk
    arrow(5.2, 3.55, 5.8, 2.98+0.12)       # Next.js → Claude
    arrow(5.2, 3.35, 5.8, 1.98+0.12)       # Next.js → Stripe
    arrow(7.4, 3.7, 7.8, 3.7)             # Next.js → Neon
    arrow(7.4, 3.3, 7.8, 2.35)            # Next.js → Redis

    # Legend
    ax.text(0.2, 0.7, "Legend:", fontsize=7, color="#475569", fontweight="bold")
    for i, (col, lbl) in enumerate([
        (C_BLUE, "Frontend"), (C_GREEN, "Backend / Vercel"),
        (C_PURPLE, "Auth"), (C_ORANGE, "AI"), ("#E11D48", "Payments"), (C_TEAL, "Database"),
    ]):
        xi = 0.2 + i * 1.6
        ax.add_patch(mpatches.Rectangle((xi, 0.35), 0.25, 0.18, color=col, zorder=3))
        ax.text(xi + 0.32, 0.44, lbl, fontsize=6.5, color="#334155", va="center")

    return _save_fig_to_bytes(fig)


def make_user_journey_diagram() -> bytes | None:
    if not HAS_MPL:
        return None
    fig, ax = plt.subplots(figsize=(10, 3.5))
    ax.set_xlim(0, 10); ax.set_ylim(0, 3.5); ax.axis("off")
    fig.patch.set_facecolor("#F8FAFC")
    ax.set_facecolor("#F8FAFC")

    C_NAVY = "#000936"; C_GOLD = "#FDC800"; C_BLUE = "#0066CB"
    C_GREEN = "#059669"; C_ORANGE = "#E34C00"

    ax.text(5, 3.3, "SelectEd — User Journey", ha="center", fontsize=12, fontweight="bold", color=C_NAVY)

    stages = [
        ("Guest",       "20 Q/day\nNo sign-up",        "#EFF6FF", C_BLUE),
        ("Sign Up",     "7-day free trial\nFull access", "#F0FDF4", C_GREEN),
        ("Free Trial",  "$0 for 7 days\nAll features",  "#FFF7ED", C_ORANGE),
        ("Premium",     "$9.99/mo AUD\nUnlimited",      "#FFFBEB", "#D97706"),
        ("Family Plan", "Up to 5 profiles\nDashboard",  "#FAF5FF", "#7C3AED"),
    ]

    for i, (title, sub, fill, edge) in enumerate(stages):
        x = 0.4 + i * 1.95
        rect = FancyBboxPatch((x, 0.6), 1.65, 1.9,
                               boxstyle="round,pad=0.1",
                               facecolor=fill, edgecolor=edge, linewidth=2, zorder=2)
        ax.add_patch(rect)
        ax.text(x + 0.825, 1.9, title, ha="center", va="center",
                fontsize=8.5, fontweight="bold", color=edge, zorder=3)
        ax.text(x + 0.825, 1.3, sub, ha="center", va="center",
                fontsize=7, color="#475569", zorder=3)

        if i < len(stages) - 1:
            ax.annotate("", xy=(x + 1.75, 1.55), xytext=(x + 1.65, 1.55),
                        arrowprops=dict(arrowstyle="-|>", color="#94A3B8", lw=1.8), zorder=4)

        # % conversion labels
        pcts = ["", "~60% sign up", "~80% continue", "~40% convert", "~25% expand"]
        if pcts[i]:
            ax.text(x + 1.73, 1.75, pcts[i], ha="center", fontsize=5.5, color="#94A3B8", style="italic")

    ax.text(5, 0.2, "Conversion funnel — illustrative figures", ha="center", fontsize=6.5, color="#94A3B8", style="italic")
    return _save_fig_to_bytes(fig)


def make_learning_modes_diagram() -> bytes | None:
    if not HAS_MPL:
        return None
    fig, ax = plt.subplots(figsize=(8, 5))
    ax.set_xlim(0, 8); ax.set_ylim(0, 5); ax.axis("off")
    fig.patch.set_facecolor("#F8FAFC")
    ax.set_facecolor("#F8FAFC")

    C_NAVY = "#000936"
    ax.text(4, 4.7, "Four Learning Modes", ha="center", fontsize=12, fontweight="bold", color=C_NAVY)

    modes = [
        (0.2, 2.55, "AI Chat Tutor",  "Socratic method\nUnlimited questions\nStep-by-step guidance", "#EFF6FF", "#0066CB"),
        (4.1, 2.55, "Practice Mode",  "Exam-style questions\nUp to 5 attempts\nFull solution reveal",  "#F0FDF4", "#059669"),
        (0.2, 0.2,  "Mock Exams",     "Timed assessment\nInstant grading\nAI mistake walkthrough",     "#FFF7ED", "#E34C00"),
        (4.1, 0.2,  "Adventure Mode", "Gamified worlds\nXP & combo system\nProgressive unlocking",     "#F5F3FF", "#7C3AED"),
    ]

    for (x, y, title, desc, fill, edge) in modes:
        rect = FancyBboxPatch((x, y), 3.7, 2.1, boxstyle="round,pad=0.12",
                               facecolor=fill, edgecolor=edge, linewidth=2.2, zorder=2)
        ax.add_patch(rect)
        ax.text(x + 1.85, y + 1.72, title, ha="center", fontsize=10, fontweight="bold", color=edge, va="center", zorder=3)
        ax.text(x + 1.85, y + 1.0, desc, ha="center", va="center", fontsize=7.5,
                color="#475569", zorder=3, linespacing=1.5)

    return _save_fig_to_bytes(fig)


def insert_diagram(doc: Document, img_bytes: bytes | None, caption: str, width: float = 6.0):
    if img_bytes is None:
        return
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_spacing(p_img, 60, 20)
    run = p_img.add_run()
    run.add_picture(io.BytesIO(img_bytes), width=Inches(width))

    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_spacing(p_cap, 0, 80)
    rc = p_cap.add_run(caption)
    rc.font.name = "Calibri"; rc.font.size = Pt(8.5)
    rc.font.italic = True; rc.font.color.rgb = MID_GRAY


def main():
    if not SRC.exists():
        print(f"ERROR: {SRC} not found"); sys.exit(1)

    if not LOGO_PNG.exists():
        print("WARNING: scripts/logo.png missing -- run: node scripts/svg_to_png.js")

    if DEST.exists():
        DEST.unlink()
        print(f"Removed old {DEST.name}")

    md_text = SRC.read_text(encoding="utf-8")
    doc     = build_doc(md_text)

    # Append visual diagrams appendix
    if HAS_MPL:
        doc.add_page_break()
        h = doc.add_paragraph(style="Heading 1")
        h.add_run("Visual Architecture Diagrams")

        doc.add_paragraph(style="Heading 2").add_run("System Architecture")
        insert_diagram(doc, make_architecture_diagram(),
                       "Figure 1 — SelectEd system architecture overview")

        doc.add_paragraph(style="Heading 2").add_run("User Journey & Conversion Funnel")
        insert_diagram(doc, make_user_journey_diagram(),
                       "Figure 2 — User journey from guest to premium family plan")

        doc.add_paragraph(style="Heading 2").add_run("Four Learning Modes")
        insert_diagram(doc, make_learning_modes_diagram(),
                       "Figure 3 — Overview of the four learning modes")
        print("  Diagrams appended (matplotlib)")
    else:
        print("  Skipped diagrams (matplotlib not installed)")

    doc.save(str(DEST))
    print(f"Saved {DEST.name}  ({DEST.stat().st_size // 1024} KB)")


if __name__ == "__main__":
    main()
